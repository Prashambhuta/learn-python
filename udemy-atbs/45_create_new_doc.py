#! usr/bin/env python3

"""
Create a new doc document.
"""

import docx

# create a new blank document object
new_doc_object = docx.Document()

# add paragraph
new_doc_object.add_paragraph("Hello")

# create a title
par_object = new_doc_object.paragraphs

# adding text to the document
par_object[0].text = "This is a new document"

# adding style
par_object[0].style = "Title"

# new paragraph
new_doc_object.add_paragraph("This is the subtitle")
par_object = new_doc_object.paragraphs
par_object[1].style = "Subtitle"

# normal body paragraph
new_doc_object.add_paragraph("""
Hello, this document is generated using python script. 
This is the body text of the document, okay, bye.
""")

# adding attribute to Subtitle para
new_doc_object.paragraphs[1].add_run(' This is the second run.')
new_doc_object.paragraphs[1].runs[1].bold = True

# saving the document
new_doc_object.save('other/new_document.docx')
