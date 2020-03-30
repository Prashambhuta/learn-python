#! usr/bin/env python3

"""
Working with word (.docx) documents
"""

# install the `python-docx` module
# download the demo file from `http://autbor.com/demo.docx`

# import the `docx` module
import docx

# create the `document object` with the demo document in `other` folder
doc_object = docx.Document('other/demo.docx')

# doc_object has paragraph member which contains the paragraph objects
print(doc_object.paragraphs)
par_object = doc_object.paragraphs

# each paragraphs object has text member objects containing the strings.
for text_object in par_object:
    print(text_object.text)


# each par_object has run member which contains the run `string` members.
# `run` is when the format of the text in par_obj changes
    run_object = text_object.runs
    for i in range(len(run_object)):
        print(run_object[i].text)

# each run member has attribute variable (italic, bold, underline) which are
# either set to True or False.
        if run_object[i].bold:
            print("This ^ is bold text.")
            # changing the above to bold +italic + underline
            run_object[i].italic = True
            run_object[i].underline = True
            run_object[i].text = 'bold + italic + underlined'

# save the document
doc_object.save('other/demo.docx')
